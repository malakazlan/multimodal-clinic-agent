"""
Document Processor for the Healthcare Voice AI Assistant.
Handles document parsing, text extraction, and chunking for different file formats.
"""

import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio

from loguru import logger

from config.settings import get_settings


class DocumentProcessor:
    """Processes documents of various formats and chunks them for RAG processing."""
    
    def __init__(self):
        self.settings = get_settings()
        self.supported_formats = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.html': self._process_html
        }
        
        logger.info("Document processor initialized")
    
    async def process_document(self, file_path: str) -> List[str]:
        """
        Process a document file and return text chunks.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of text chunks
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Document file not found: {file_path}")
            
            # Check file format
            file_ext = file_path.suffix.lower()
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            logger.info(f"Processing document: {file_path.name} ({file_ext})")
            
            # Extract text based on file format
            processor = self.supported_formats[file_ext]
            text = await processor(file_path)
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from document: {file_path.name}")
                return []
            
            # Clean and preprocess text
            cleaned_text = self._clean_text(text)
            
            # Chunk the text
            chunks = self.chunk_text(
                text=cleaned_text,
                chunk_size=self.settings.chunk_size,
                chunk_overlap=self.settings.chunk_overlap
            )
            
            logger.info(f"Processed document: {file_path.name} -> {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {str(e)}", exc_info=True)
            raise
    
    async def _process_text(self, file_path: Path) -> str:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Process Markdown files."""
        try:
            import markdown
            
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert markdown to HTML
            html = markdown.markdown(md_content)
            
            # Extract text from HTML
            return self._extract_text_from_html(html)
            
        except ImportError:
            logger.warning("markdown library not available, processing as plain text")
            return await self._process_text(file_path)
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files."""
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            return text
            
        except ImportError:
            logger.warning("PyPDF2 library not available, cannot process PDF")
            return ""
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {str(e)}")
            return ""
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process DOCX files."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text
            
        except ImportError:
            logger.warning("python-docx library not available, cannot process DOCX")
            return ""
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {str(e)}")
            return ""
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML files."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            return self._extract_text_from_html(html_content)
            
        except ImportError:
            logger.warning("beautifulsoup4 library not available, processing as plain text")
            return await self._process_text(file_path)
        except Exception as e:
            logger.error(f"Failed to process HTML {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_html(self, html_content: str) -> str:
        """Extract clean text from HTML content."""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract text from HTML: {str(e)}")
            return html_content
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and preprocess extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def chunk_text(
        self,
        text: str,
        chunk_size: int = None,
        chunk_overlap: int = None
    ) -> List[str]:
        """
        Split text into overlapping chunks for RAG processing.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            chunk_overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        # Use settings defaults if not provided
        chunk_size = chunk_size or self.settings.chunk_size
        chunk_overlap = chunk_overlap or self.settings.chunk_overlap
        
        # Ensure overlap is not larger than chunk size
        chunk_overlap = min(chunk_overlap, chunk_size - 1)
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + chunk_size
            
            # Extract chunk
            chunk = text[start:end]
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings
                sentence_endings = ['.', '!', '?', '\n']
                for ending in sentence_endings:
                    last_ending = chunk.rfind(ending)
                    if last_ending > chunk_size * 0.7:  # Only break if ending is in last 30%
                        chunk = chunk[:last_ending + 1]
                        end = start + last_ending + 1
                        break
            
            chunks.append(chunk.strip())
            
            # Move start position for next chunk
            start = end - chunk_overlap
            
            # Ensure we don't get stuck in an infinite loop
            if start >= len(text):
                break
        
        logger.debug(f"Text chunked into {len(chunks)} chunks (size: {chunk_size}, overlap: {chunk_overlap})")
        return chunks
    
    def chunk_text_by_sentences(
        self,
        text: str,
        max_chunk_size: int = None
    ) -> List[str]:
        """
        Split text into chunks based on sentence boundaries.
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum size of each chunk
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        max_chunk_size = max_chunk_size or self.settings.chunk_size
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) + 1 > max_chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        logger.debug(f"Text chunked into {len(chunks)} sentence-based chunks")
        return chunks
    
    def get_document_stats(self, text: str) -> Dict[str, Any]:
        """
        Get statistics about a document.
        
        Args:
            text: Document text
            
        Returns:
            Dictionary with document statistics
        """
        if not text:
            return {
                "characters": 0,
                "words": 0,
                "sentences": 0,
                "paragraphs": 0,
                "estimated_tokens": 0
            }
        
        # Basic statistics
        char_count = len(text)
        word_count = len(text.split())
        sentence_count = len(re.split(r'[.!?]+', text))
        paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
        
        # Estimate token count (rough approximation: 4 characters per token)
        estimated_tokens = char_count // 4
        
        stats = {
            "characters": char_count,
            "words": word_count,
            "sentences": sentence_count,
            "paragraphs": paragraph_count,
            "estimated_tokens": estimated_tokens
        }
        
        return stats
    
    async def health_check(self) -> bool:
        """
        Perform health check on the document processor.
        
        Returns:
            True if healthy, False otherwise
        """
        try:
            # Test text processing
            test_text = "This is a test document. It contains multiple sentences. We can process it."
            chunks = self.chunk_text(test_text, chunk_size=50, chunk_overlap=10)
            
            if not chunks or len(chunks) == 0:
                return False
            
            # Test text cleaning
            dirty_text = "  This   has   excessive   whitespace  \n\n\n"
            cleaned = self._clean_text(dirty_text)
            
            if not cleaned or len(cleaned) == 0:
                return False
            
            logger.debug("Document processor health check passed")
            return True
            
        except Exception as e:
            logger.error(f"Document processor health check failed: {str(e)}", exc_info=True)
            return False
